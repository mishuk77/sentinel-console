from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colours ──────────────────────────────────────────────────────────────────
COL_BLUE    = RGBColor(0x1e, 0x40, 0xab)
COL_WHITE   = RGBColor(0xff, 0xff, 0xff)
COL_DARK    = RGBColor(0x0f, 0x17, 0x2a)
COL_RED     = RGBColor(0xdc, 0x26, 0x26)
COL_GREEN   = RGBColor(0x16, 0xa3, 0x4a)
COL_MUTED   = RGBColor(0x64, 0x74, 0x8b)
COL_USER    = RGBColor(0x1d, 0x4e, 0xd8)
COL_ORANGE  = RGBColor(0xea, 0x58, 0x0c)

HEX_BLUE    = '1e40ab'
HEX_SLATE   = '334155'
HEX_LIGHT   = 'f1f5f9'
HEX_USER_BG = 'eff6ff'
HEX_GREEN_BG= 'f0fdf4'
HEX_AMBER   = 'fef9c3'

ALGO_DISPLAY = {
    'logistic_regression': 'Logistic Regression (ElasticNet / L2)',
    'random_forest':       'Random Forest Classifier',
    'xgboost':             'Extreme Gradient Boosting Classifier (XGBoost)',
}

ALGO_DESC = {
    'logistic_regression': {
        'intro': (
            "Logistic regression models the log-odds of the binary target as a linear combination of "
            "input features. It is a class of generalised linear models (GLM) that uses the binomial "
            "distribution to fit regression models to a binary (0/1) response variable. "
            "L2 (Ridge) regularisation prevents overfitting by penalising large coefficients; balanced "
            "class weights compensate for target imbalance. Optimisation uses the limited-memory "
            "Broyden–Fletcher–Goldfarb–Shanno (L-BFGS) algorithm with up to 1,000 iterations."
        ),
        'extra': (
            "A practical feature of regularised logistic regression is that it provides coefficient-level "
            "interpretability: each coefficient directly indicates the direction and relative magnitude "
            "of each feature's effect on the log-odds of the outcome, making it straightforward to "
            "produce SR-11-7 compliant rationale for individual model decisions."
        ),
        'params': [
            ['Type',    'Parameter',       'Description',                                            'Value'],
            ['select',  'Regularisation',  'Penalty term applied to coefficient magnitudes',          'L2 (Ridge)'],
            ['float',   'C',               'Inverse regularisation strength (smaller = stronger reg)','1.0'],
            ['select',  'Class Weights',   'Reweights classes to handle imbalance',                   'Balanced (auto)'],
            ['select',  'Solver',          'Optimisation algorithm',                                  'lbfgs'],
            ['int',     'Max Iterations',  'Maximum number of solver iterations',                     '1,000'],
            ['bool',    'Fit Intercept',   'Whether to include a bias/intercept term',               'True'],
        ],
    },
    'random_forest': {
        'intro': (
            "Random Forest is an ensemble learning method that constructs a large number of decision "
            "trees on bootstrapped training subsets (bagging) and averages their probability outputs. "
            "Feature subsampling (sqrt of total features) at each split reduces inter-tree correlation, "
            "improving generalisation over a single decision tree. Balanced class weights are applied "
            "independently at each tree to handle class imbalance."
        ),
        'extra': (
            "Random Forests are non-parametric and make no distributional assumptions about input features. "
            "They are naturally robust to outliers and can capture complex non-linear interactions. "
            "Feature importance is computed via mean decrease in impurity (MDI) across all trees."
        ),
        'params': [
            ['Type',    'Parameter',       'Description',                                            'Value'],
            ['int',     'n_estimators',    'Number of trees in the forest',                          '50'],
            ['int',     'max_depth',       'Maximum depth of each tree',                             '10'],
            ['select',  'Class Weights',   'Reweights classes to handle imbalance',                   'Balanced (auto)'],
            ['select',  'Feature Sampling','Features considered at each split',                       'sqrt(n_features)'],
            ['select',  'Criterion',       'Quality measure used for splits',                         'gini'],
            ['int',     'Random State',    'Seed for reproducibility',                               '42'],
        ],
    },
    'xgboost': {
        'intro': (
            "Extreme Gradient Boosting (XGBoost) implements regularised gradient-boosted decision trees. "
            "Trees are added sequentially, each correcting the residual errors of the prior ensemble. "
            "XGBoost uses second-order Taylor approximations of the loss function during tree construction, "
            "making it significantly faster and more accurate than standard GBM implementations. "
            "scale_pos_weight = negative_count / positive_count is set automatically to handle class imbalance. "
            "LogLoss (binomial deviance) is the training objective."
        ),
        'extra': (
            "Gradient Boosting Machines (GBMs) are a generalisation of Freund and Schapire's AdaBoost "
            "algorithm (1995) that handles arbitrary differentiable loss functions. Due to their iterative "
            "nature, GBMs are prone to overfitting the training data given enough iterations. "
            "The two critical hyperparameters to control overfitting are the learning rate and the number "
            "of trees (n_estimators). XGBoost additionally applies L1 and L2 regularisation at the leaf level."
        ),
        'params': [
            ['Type',       'Parameter',         'Description',                                         'Value'],
            ['select',     'eval_metric',        'Loss function used during training',                  'logloss'],
            ['int',        'max_depth',          'Maximum depth of each boosted tree',                  '6'],
            ['float',      'learning_rate (eta)','Shrinkage factor per boosting step',                  '0.3'],
            ['select',     'scale_pos_weight',   'Balances positive and negative weights',              'neg/pos ratio (auto)'],
            ['select',     'use_label_encoder',  'Use internal label encoding',                         'False'],
            ['int',        'n_jobs',             'Parallelism (restricted to avoid OOM)',               '1'],
        ],
    },
}

ALGO_LITERATURE = {
    'logistic_regression': [
        "Hosmer Jr, David W., and Stanley Lemeshow. Applied logistic regression. John Wiley & Sons, 2004.",
        "Harrell, Frank E. Regression modeling strategies: with applications to linear models, logistic regression, and survival analysis. Springer, 2015.",
        "Tibshirani, Robert. Regression shrinkage and selection via the lasso. Journal of the Royal Statistical Society: Series B 58.1 (1996): 267–288.",
        "Hoerl, Arthur E., and Robert W. Kennard. Ridge regression: Biased estimation for nonorthogonal problems. Technometrics 12.1 (1970): 55–67.",
    ],
    'random_forest': [
        "Breiman, Leo. Random Forests. Machine Learning 45.1 (2001): 5–32.",
        "Breiman, Leo. Arcing the edge. Technical Report 486, Statistics Dept, University of California at Berkeley, 1997.",
        "Hastie, T., Tibshirani, R., and Friedman, J. Elements of Statistical Learning. Springer, 2009.",
        "Liaw, Andy, and Matthew Wiener. Classification and regression by randomForest. R news 2.3 (2002): 18–22.",
    ],
    'xgboost': [
        "Chen, Tianqi, and Carlos Guestrin. XGBoost: A scalable tree boosting system. Proceedings of the 22nd ACM SIGKDD, 2016.",
        "Friedman, Jerome H. Greedy function approximation: a gradient boosting machine. Annals of statistics (2001): 1189–1232.",
        "Freund, Yoav, and Robert E. Schapire. A decision-theoretic generalisation of on-line learning and an application to boosting. Journal of Computer and System Sciences 55.1 (1997): 119–139.",
        "Chen, T, and He, T. Higgs Boson Discovery with Boosted Trees. JMLR: Workshop and Conference Proceedings, 2015.",
    ],
}

# ── XML helpers ───────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_para_bg(para, hex_color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def _cell_run(cell, text: str, bold=False, italic=False, color=None, size=9, align=None):
    para = cell.paragraphs[0]
    para.clear()
    if align:
        para.alignment = align
    run = para.add_run(str(text) if text is not None else '—')
    run.font.bold = bold
    run.font.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

# ── Reusable builders ─────────────────────────────────────────────────────────

def _h1(doc, text):
    """Section heading (e.g. "3   Executive Summary")."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after  = Pt(6)
    run = para.add_run(text)
    run.font.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = COL_BLUE
    return para


def _h2(doc, text):
    """Subsection heading (e.g. "3.1   Model Stakeholders")."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.font.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = COL_DARK
    return para


def _h3(doc, text):
    """Sub-subsection heading (e.g. "5.3.1   XGBoost Classifier")."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(3)
    run = para.add_run(text)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COL_DARK
    return para


def _h4(doc, text):
    """Deep heading (e.g. "5.6.3.1   Model Features and Summary Statistics")."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(3)
    run = para.add_run(text)
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = COL_MUTED
    return para


def _body(doc, text: str):
    para = doc.add_paragraph(text)
    for run in para.runs:
        run.font.size = Pt(9)
    para.paragraph_format.space_after = Pt(4)
    return para


def _caption(doc, text: str):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = COL_MUTED
    para.paragraph_format.space_after = Pt(3)
    return para


def _user_fill(doc, text: str):
    """Blue-shaded block indicating a section the Model Owner must complete."""
    para = doc.add_paragraph()
    label = para.add_run("[REQUIRED: User Input]  ")
    label.font.bold   = True
    label.font.size   = Pt(9)
    label.font.color.rgb = COL_USER
    body = para.add_run(text)
    body.font.size    = Pt(9)
    body.font.italic  = True
    body.font.color.rgb = COL_USER
    _set_para_bg(para, HEX_USER_BG)
    para.paragraph_format.left_indent  = Inches(0.15)
    para.paragraph_format.right_indent = Inches(0.15)
    para.paragraph_format.space_after  = Pt(6)
    para.paragraph_format.space_before = Pt(4)
    return para


def _quote(doc, text: str):
    """Italic block-quote style (like DataRobot's intense_quote paragraphs)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size   = Pt(9)
    run.font.italic = True
    run.font.color.rgb = COL_MUTED
    para.paragraph_format.left_indent  = Inches(0.3)
    para.paragraph_format.right_indent = Inches(0.3)
    para.paragraph_format.space_after  = Pt(4)
    return para


def _bullet(doc, text: str, level: int = 1):
    style = 'List Bullet' if level == 1 else f'List Bullet {level}'
    try:
        para = doc.add_paragraph(style=style)
    except Exception:
        para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    run.font.size = Pt(9)
    return para


def _spacer(doc, pts=4):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(pts)
    return para


def _kv_table(doc, rows, col_widths=(1.9, 4.6)):
    """Two-column key/value table."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    for ri, (key, val) in enumerate(rows):
        bg = HEX_LIGHT if ri % 2 == 0 else None
        _cell_run(table.rows[ri].cells[0], key, bold=True, color=COL_BLUE)
        _cell_run(table.rows[ri].cells[1], val)
        if bg:
            _set_cell_bg(table.rows[ri].cells[0], bg)
            _set_cell_bg(table.rows[ri].cells[1], bg)
    for cell in table.columns[0].cells:
        cell.width = Inches(col_widths[0])
    for cell in table.columns[1].cells:
        cell.width = Inches(col_widths[1])
    return table


def _data_table(doc, headers, rows, col_widths_in, header_bg=HEX_BLUE):
    """Full table with coloured header + alternating row shading."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    for ci, h in enumerate(headers):
        _cell_run(table.rows[0].cells[ci], h, bold=True, color=COL_WHITE)
        _set_cell_bg(table.rows[0].cells[ci], header_bg)
    for ri, row_data in enumerate(rows):
        bg = HEX_LIGHT if ri % 2 == 0 else None
        for ci, val in enumerate(row_data):
            _cell_run(table.rows[ri + 1].cells[ci], val)
            if bg:
                _set_cell_bg(table.rows[ri + 1].cells[ci], bg)
    for ci, w in enumerate(col_widths_in):
        for cell in table.columns[ci].cells:
            cell.width = Inches(w)
    return table


# ── Main entry point ──────────────────────────────────────────────────────────

def generate_model_documentation(model, sibling_models=None) -> BytesIO:
    doc = Document()

    for section in doc.sections:
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    metrics               = model.metrics or {}
    auc                   = metrics.get('auc') or 0.0
    gini                  = metrics.get('gini') if metrics.get('gini') is not None else (2 * auc - 1)
    cv_fold_scores        = metrics.get('cv_fold_scores') or []
    cv_auc_mean           = metrics.get('cv_auc_mean')
    cv_auc_std            = metrics.get('cv_auc_std')
    feature_importance    = metrics.get('feature_importance') or []
    calibration           = metrics.get('calibration') or []
    feature_stats         = metrics.get('feature_stats') or []
    classification_metrics= metrics.get('classification_metrics') or {}
    data_profile          = metrics.get('data_profile') or {}

    algo         = model.algorithm or ''
    algo_display = ALGO_DISPLAY.get(algo, algo.replace('_', ' ').title())
    created_str  = model.created_at.strftime('%Y-%m-%d %H:%M UTC') if model.created_at else '—'
    today_str    = datetime.utcnow().strftime('%Y-%m-%d')
    now_str      = datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')
    status_str   = str(model.status.value if hasattr(model.status, 'value') else model.status)

    def _fmt_pct(v, decimals=4):
        return f"{v:.{decimals}f}" if v is not None else '—'

    def _assess(v):
        return 'Strong' if v > 0.80 else ('Acceptable' if v > 0.75 else 'Below Benchmark')

    # ── COVER PAGE ────────────────────────────────────────────────────────────
    title_p = doc.add_heading('SENTINEL', 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_p.runs:
        run.font.color.rgb = COL_BLUE
        run.font.size = Pt(32)

    sub_p = doc.add_paragraph('Model Development Documentation')
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub_p.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = COL_MUTED

    _spacer(doc, 8)

    name_p = doc.add_paragraph(model.name or 'Unnamed Model')
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in name_p.runs:
        run.font.bold = True
        run.font.size = Pt(20)

    algo_p = doc.add_paragraph(algo_display)
    algo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in algo_p.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = COL_MUTED

    _spacer(doc, 8)
    _kv_table(doc, [
        ('Model ID',         model.id or '—'),
        ('Algorithm',        algo_display),
        ('Status',           status_str),
        ('Training Date',    created_str),
        ('AUC (Hold-out)',   f"{auc:.4f}  ({auc * 100:.2f}%)"),
        ('Gini Coefficient', f"{gini:.4f}  ({gini * 100:.1f}%)"),
        ('Prepared By',      'Sentinel AI — Automated Documentation'),
        ('Date Generated',   today_str),
    ], col_widths=(2.0, 4.5))

    _spacer(doc, 8)
    note_p = doc.add_paragraph()
    nr = note_p.add_run(
        "This document was automatically generated by Sentinel and follows regulatory model documentation "
        "guidance (FRB SR-11-7, OCC 2011-12). Sections marked [REQUIRED: User Input] must be completed "
        "by the Model Owner prior to submission to the Model Risk Management function. "
        "Copyright \u00a9 Sentinel AI.")
    nr.font.size = Pt(8)
    nr.font.italic = True
    nr.font.color.rgb = COL_MUTED

    doc.add_page_break()

    # ── SECTION 1: HOW TO USE THIS DOCUMENT ──────────────────────────────────
    _h1(doc, "1   How To Use This Document")
    _body(doc,
        "Highly regulated industries, such as banking and insurance, must comply with government "
        "regulations for model validation before a model can be put into production. Sentinel's model "
        "documentation is designed to assist organisations in meeting these regulatory requirements by "
        "providing automatically generated, evidence-based content for each trained model.")
    _body(doc,
        "This document is split into two components: those sections that are automatically produced by "
        "Sentinel and those that require further input by the user (Model Owner). Sections requiring user "
        "input are clearly marked with a blue [REQUIRED: User Input] label. These sections must be "
        "completed before the document is submitted to the Model Risk Management function.")
    _body(doc, f"Copyright \u00a9 {datetime.utcnow().year}, Sentinel AI, Inc.")

    _spacer(doc, 8)

    # ── TABLE OF CONTENTS ─────────────────────────────────────────────────────
    _h1(doc, "Table of Contents")
    toc_entries = [
        (1, "1   How To Use This Document"),
        (1, "2   Sentinel Model Development Documentation"),
        (1, "3   Executive Summary and Model Overview"),
        (2, "3.1   Model Stakeholders"),
        (2, "3.2   Model Development Purpose and Intended Use"),
        (2, "3.3   Model Description and Overview"),
        (2, "3.4   Overview of Model Results"),
        (2, "3.5   Model Interdependencies"),
        (1, "4   Model Data Overview"),
        (2, "4.1   Feature Association"),
        (2, "4.2   Data Source Overview and Appropriateness"),
        (2, "4.3   Input Data Extraction, Preparation, and Quality & Completeness"),
        (2, "4.4   Data Assumptions"),
        (1, "5   Model Theoretical Framework and Methodology"),
        (2, "5.1   Model Development Overview"),
        (2, "5.2   Model Assumptions"),
        (2, "5.3   Model Methodology"),
        (3, f"5.3.1   {algo_display}"),
        (2, "5.4   Literature Review and References"),
        (2, "5.5   Alternative Model Frameworks and Theories Considered"),
        (2, "5.6   Variable Selection"),
        (3, "5.6.1   Sentinel Quantitative Analysis"),
        (3, "5.6.2   Expert Judgement and Variable Selection"),
        (3, "5.6.3   Final Model Variables"),
        (4, "5.6.3.1   Model Features and Summary Statistics"),
        (1, "6   Model Performance and Stability"),
        (2, "6.1   Model Validation Stability"),
        (3, "6.1.1   Cross Validation Scores"),
        (3, "6.1.2   Data Partitioning Methodology"),
        (2, "6.2   Model Performance (Sample Scores)"),
        (2, "6.3   Sensitivity Testing and Analysis"),
        (3, "6.3.1   Lift Chart"),
        (3, "6.3.2   Key Relationships"),
        (3, "6.3.3   Sensitivity Analysis"),
        (3, "6.3.4   Accuracy (Receiver Operating Characteristic)"),
        (1, "7   Model Implementation and Output Reporting"),
        (2, "7.1   Version Control"),
        (2, "7.2   Monitoring and Ongoing Validation"),
    ]
    for level, entry in toc_entries:
        style = 'List Bullet' if level == 1 else f'List Bullet {min(level, 3)}'
        try:
            p = doc.add_paragraph(style=style)
        except Exception:
            p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(entry)
        run.font.size = Pt(9)
        if level == 1:
            run.font.bold = True

    doc.add_page_break()

    # ── SECTION 2: SENTINEL MODEL DEVELOPMENT DOCUMENTATION ──────────────────
    _h1(doc, "2   Sentinel Model Development Documentation")
    for q in [
        ("A key component of effective model risk management is sufficiently detailed documentation for "
         "model development, implementation, and use, so that parties other than the original "
         "developers can use and validate the model.  — FRB SR-11-7"),
        ("The purpose of this document is not to be prescriptive in format and content, but rather to "
         "serve as a guide in creating sufficiently rigorous model documentation. This guide should be "
         "viewed in the spirit of its intent — that of thorough, transparent, and explainable "
         "documentation of model development and performance.  — OCC 2011-12"),
    ]:
        _quote(doc, q)

    doc.add_page_break()

    # ── SECTION 3: EXECUTIVE SUMMARY ─────────────────────────────────────────
    _h1(doc, "3   Executive Summary and Model Overview")

    _h2(doc, "3.1   Model Stakeholders")
    _user_fill(doc,
        "Describe the model's purpose and its intended business use. Describe all stakeholders of this "
        "model, including their role, line-of-business, and team.\n\n"
        "Model Owner(s): The individual who owns the business risk addressed by the model and provides "
        "approval for the model to be used within the line-of-business.\n\n"
        "Model Developer(s): The individual responsible for building new models with Sentinel or "
        "maintaining existing models.\n\n"
        "Model User(s): Those teams who will use the model output as part of their ongoing business "
        "operations.\n\n"
        "Model Validator(s): The validators are responsible for independent model review and approval "
        "prior to its first use.")

    _h2(doc, "3.2   Model Development Purpose and Intended Use")
    _user_fill(doc,
        "Describe the model's purpose, including a summary of the business need for this particular "
        "model. Concisely describe how the model will be used to add business value. Describe all "
        "approved use cases, any constraints on its application, and the populations to which the "
        "model will be applied. Per SR-11-7, all intended uses must be documented and approved.")

    _h2(doc, "3.3   Model Description and Overview")
    _body(doc,
        f"The particular model referenced in this document is {model.name or 'Unnamed Model'}. "
        f"This model was developed using the Sentinel AI automated machine learning platform. "
        f"The model was trained on {created_str} and assigned Model ID: {model.id}.")
    _body(doc, "The model development workflow includes the following elements:")
    _bullet(doc, algo_display)
    _spacer(doc)
    _kv_table(doc, [
        ('Model Name',         model.name or '—'),
        ('Model ID',           model.id or '—'),
        ('Algorithm',          algo_display),
        ('Status',             status_str),
        ('Training Date',      created_str),
        ('Artifact Path',      model.artifact_path or '—'),
        ('Decision System ID', model.decision_system_id or '—'),
    ])

    _spacer(doc)
    _h2(doc, "3.4   Overview of Model Results")
    _body(doc,
        "Sentinel runs performance testing during the model development process to evaluate model "
        "results and reliability. The cross-validation and hold-out scores are presented below.")
    _caption(doc, "Table 1 — Model Performance Scores")

    score_rows = []
    if cv_auc_mean is not None:
        score_rows.append(['cross_validation', _fmt_pct(cv_auc_mean)])
    score_rows.append(['holdout', _fmt_pct(auc)])
    if cv_auc_std is not None:
        score_rows.append(['cv_std_dev', f"± {_fmt_pct(cv_auc_std)}"])

    _data_table(doc,
        ['Scoring Type', 'AUC Score'],
        score_rows,
        [2.5, 2.5])

    _spacer(doc)
    _h2(doc, "3.5   Model Interdependencies")
    _user_fill(doc,
        "Understanding interdependent relationships allows for enhanced understanding and improved "
        "ability to manage and aggregate model risk. Explain how this model interacts with other models "
        "or systems within the risk decisioning framework, including any upstream data feeds or "
        "downstream models that consume this model's output scores.")

    doc.add_page_break()

    # ── SECTION 4: DATA OVERVIEW ──────────────────────────────────────────────
    _h1(doc, "4   Model Data Overview")

    _h2(doc, "4.1   Feature Association")
    _body(doc,
        "Sentinel computes feature-level statistics during training to identify informative features, "
        "detect potential target leakage, and characterise the modelling dataset. The Feature Summary "
        "Statistics table in Section 5.6.3.1 provides a quantitative overview of each model feature, "
        "including data type, missing rate, and distributional statistics.")
    _body(doc, "Some key benefits of feature analysis include:")
    for b in [
        "Understand the strength and nature of associations between features and the target;",
        "Detect potential target leakage based on correlation thresholds;",
        "Identify high-cardinality or near-constant features prior to model building.",
    ]:
        _bullet(doc, b)
    if data_profile:
        _spacer(doc)
        _kv_table(doc, [
            ('Total Rows in Dataset',  f"{data_profile.get('total_rows', 0):,}"),
            ('Feature Count',          str(data_profile.get('feature_count', 0))),
            ('Overall Missing Rate',   f"{data_profile.get('missing_pct', 0):.2f}%"),
            ('Target Default Rate',    f"{data_profile.get('class_balance', 0) * 100:.2f}%"),
        ])

    _h2(doc, "4.2   Data Source Overview and Appropriateness")
    _user_fill(doc,
        "Explain how the data is suitable and relevant for the business problem and model use. Describe "
        "how, and from where, the data was obtained. Provide a detailed description of the data source "
        "and its relevance to the business problem being addressed by this model. Assess whether the "
        "data used for model development is appropriate given the populations to which the model will "
        "be applied. If the model development and model implementation data sources differ, provide a "
        "detailed explanation justifying the use of different data sources.")

    _h2(doc, "4.3   Input Data Extraction, Preparation, and Quality & Completeness")
    _user_fill(doc,
        "Provide a detailed description of the data extraction and preparation process, and discuss "
        "any analysis conducted to confirm the data are complete and of sufficient quality. Review and "
        "comment on any data weaknesses and limitations and their probable potential effects on the "
        "model. For example, data truncation, extraction errors, survivorship bias, or data drift "
        "between training and scoring populations.")

    _h2(doc, "4.4   Data Assumptions")
    _user_fill(doc,
        "Comment on data assumptions, the potential effects on the model, and any mitigating data "
        "controls. For example, assumptions related to data truncation, missing value treatment, "
        "temporal stability of features, and representativeness of the training population relative "
        "to the intended deployment population.")

    doc.add_page_break()

    # ── SECTION 5: METHODOLOGY ────────────────────────────────────────────────
    _h1(doc, "5   Model Theoretical Framework and Methodology")

    _h2(doc, "5.1   Model Development Overview")
    _body(doc,
        "Sentinel simplifies model development by performing a parallel search across multiple "
        "algorithm families, based on both the characteristics of the training data and established "
        "best practices for binary classification of credit risk events.")
    _body(doc, "The fundamental workflow within Sentinel for model development is as follows:")
    for step in [
        "Data Ingestion: User uploads a modelling dataset that includes the prediction target and selects "
        "relevant feature columns.",
        "Target Selection: User selects the prediction target; Sentinel trains all candidate models "
        "against the specified binary target variable.",
        "Automated Data Preparation: Sentinel analyses the input data and automatically performs "
        "preprocessing steps described in detail in Section 5.3, including one-hot encoding, missing "
        "value imputation, and class imbalance correction.",
        "Sentinel trains three candidate model blueprints in parallel. Each model is evaluated using "
        "5-fold stratified cross-validation on the training partition and a held-out test set.",
        "Transparent Model Evaluation and Selection: Sentinel presents all candidate models on a "
        "Leaderboard ranked by hold-out AUC. The user selects the final model based on performance "
        "and interpretability.",
        "Model Deployment and Monitoring: Once the final model is activated within a Decision System, "
        "Sentinel applies it to incoming scoring requests via the real-time decisioning API.",
    ]:
        _bullet(doc, step)

    _h2(doc, "5.2   Model Assumptions")
    _user_fill(doc,
        "This section should include model limitations, potential effects, and any mitigating controls "
        "in place. Limitations come in part from weaknesses in the data and in part from inherent "
        "characteristics of the model type.")
    _body(doc,
        "Machine learning methods can produce more accurate predictive models than traditional "
        "statistical regression methods because they are more flexible and can capture non-linear "
        "patterns in the data. However, this flexibility comes with trade-offs:")
    for a in [
        "Training data is assumed to be representative of the future scoring population (temporal "
        "stability assumption). Model performance may degrade if the population shifts over time.",
        "The target variable is assumed to accurately capture the modelled risk event. Any "
        "misclassification in the target label will reduce model quality.",
        "Feature relationships observed in training are assumed to persist in the deployment environment.",
        "Overfitting is mitigated through stratified 5-fold cross-validation and a held-out test "
        "partition (80/20 stratified split). Low variance in CV fold scores indicates stable "
        "generalisation.",
        "No distributional assumptions are required for tree-based models (XGBoost, Random Forest). "
        "Logistic Regression assumes approximately log-linear relationships between features and "
        "log-odds of the target.",
        "Missing values are imputed with the column median (numeric) fitted on the training set only, "
        "to prevent target leakage into preprocessing statistics.",
    ]:
        _bullet(doc, a)

    _h2(doc, "5.3   Model Methodology")
    _body(doc,
        "The modelling workflow consists of the following elements, which connect to create the "
        "model blueprint:")
    _bullet(doc, algo_display)
    _body(doc, "The following subsection includes details for the modelling algorithm.")

    _h3(doc, f"5.3.1   {algo_display}")
    if algo in ALGO_DESC:
        ad = ALGO_DESC[algo]
        _body(doc, ad['intro'])
        _spacer(doc)
        _body(doc, ad['extra'])
        _spacer(doc)
        param_rows = ad['params']
        _caption(doc, f"Table — {algo_display} Parameters")
        _data_table(doc, param_rows[0], param_rows[1:],
                    [0.7, 1.8, 3.3, 0.9], header_bg=HEX_SLATE)
    _spacer(doc)

    _h2(doc, "5.4   Literature Review and References")
    refs = ALGO_LITERATURE.get(algo, [])
    for ref in refs:
        _bullet(doc, ref)
    if not refs:
        _user_fill(doc, "Provide academic and industry references relevant to the algorithm and domain.")

    _h2(doc, "5.5   Alternative Model Frameworks and Theories Considered")
    _body(doc,
        "As stated by regulatory guidance, comparison with alternative theories and approaches provides "
        "guidance for final model selection and is a fundamental component of model documentation.")
    _body(doc,
        "Sentinel develops three alternative model blueprints simultaneously, exposes the details of "
        "how these models were built and how they perform, and enables the user to select the final "
        "model based on performance and interpretability. All models are trained on the same dataset "
        "with identical preprocessing.")
    _body(doc,
        "The performance metric used for this project was AUC (Area Under the Receiver Operating "
        "Characteristic Curve). The model types considered during the model selection process included "
        "the following models, ranked in order of hold-out performance:")
    _spacer(doc)
    _caption(doc, "Table 2 — Candidate Model Leaderboard (sorted by AUC descending). \u2605 = this model.")

    if sibling_models:
        sorted_sibs = sorted(sibling_models,
                             key=lambda m: (m.metrics or {}).get('auc', 0), reverse=True)
        sib_rows = []
        highlight_indices = []
        for i, sib in enumerate(sorted_sibs):
            sm    = sib.metrics or {}
            s_auc = sm.get('auc') or 0
            s_cv  = sm.get('cv_auc_mean')
            s_std = sm.get('cv_auc_std')
            is_me = sib.id == model.id
            if is_me:
                highlight_indices.append(i)
            sib_rows.append([
                ('\u2605 ' if is_me else '') + (sib.name or '—'),
                ALGO_DISPLAY.get(sib.algorithm, sib.algorithm or '—'),
                _fmt_pct(s_cv)  if s_cv  else '—',
                _fmt_pct(s_auc) if s_auc else '—',
                _fmt_pct(s_std) if s_std else '—',
                str(sib.status.value if hasattr(sib.status, 'value') else sib.status),
            ])
        table = _data_table(doc,
            ['Model Name', 'Algorithm', 'Cross Validation AUC', 'Holdout AUC', 'CV Std Dev', 'Status'],
            sib_rows,
            [1.6, 1.7, 1.3, 1.0, 0.85, 0.85])
        for idx in highlight_indices:
            for cell in table.rows[idx + 1].cells:
                _set_cell_bg(cell, HEX_GREEN_BG)
    else:
        _body(doc, "No sibling models available for comparison.")

    _h2(doc, "5.6   Variable Selection")
    _body(doc,
        "The model's variable selection process includes a balance of quantitative analysis and key "
        "domain knowledge about the underlying business problem:")
    for b in [
        "Sentinel Quantitative Analysis: Key components related to variable selection that are "
        "automated by Sentinel.",
        "Expert Judgement and Variable Selection: Summary of the expert judgement used during the "
        "variable selection process.",
        "Final Model Variables: Final feature list chosen.",
    ]:
        _bullet(doc, b)

    _h3(doc, "5.6.1   Sentinel Quantitative Analysis")
    _body(doc,
        "Sentinel automatically creates a feature list based on the columns selected by the user. "
        "During training, Sentinel excludes the following categories of columns to prevent overfitting "
        "and target leakage:")
    for b in [
        "Identifier columns: Columns with names matching common ID patterns (id, customer_id, "
        "applicant_id, uuid, etc.) are excluded automatically.",
        "High-cardinality categorical columns: String columns with more than 50 unique values are "
        "dropped as they typically represent IDs or free-text fields that cannot be meaningfully "
        "encoded without significant preprocessing.",
        "Target leakage risk is assessed using Pearson correlation (numeric features) or label-encoded "
        "correlation (categorical features) between each feature and the target. Features with "
        "|correlation| > 0.8 are flagged as High Risk; > 0.5 as Moderate Risk.",
    ]:
        _bullet(doc, b)

    _h3(doc, "5.6.2   Expert Judgement and Variable Selection")
    _user_fill(doc,
        "This section should include additional detail regarding the variable selection process and "
        "any expert judgement used during feature selection. For example: rationale for including or "
        "excluding specific features; any regulatory constraints on feature use (e.g. protected "
        "characteristics under ECOA/Reg B); and the business rationale linking each feature to the "
        "modelled risk event.")

    _h3(doc, "5.6.3   Final Model Variables")
    _body(doc,
        f"Below is the final set of model feature/independent variables used in the "
        f"{algo_display} model, as well as summary statistics.")

    _h4(doc, "5.6.3.1   Model Features and Summary Statistics")
    _body(doc,
        "The Model Features and Summary Statistics table provides a brief overview of the summary "
        "statistics of model features. This includes Feature Name, variable type, unique value count, "
        "missing value count, mean, standard deviation, median, minimum, maximum, and an assessment "
        "of target leakage risk.")
    _body(doc,
        "Target leakage risk is assessed using correlation with the target variable:\n"
        "  \u2022  High risk: |corr| > 0.8 — flagged\n"
        "  \u2022  Moderate risk: |corr| > 0.5 — flagged\n"
        "  \u2022  Low risk: |corr| \u2264 0.5 — no action")

    if feature_stats:
        _caption(doc, "Table 3 — Model Feature Summary Statistics")
        fs_rows = []
        for f in feature_stats:
            def _v(x):
                return str(x) if x is not None else '—'
            fs_rows.append([
                f.get('feature', '—'),
                f.get('var_type', '—'),
                _v(f.get('unique')),
                _v(f.get('missing')),
                _v(f.get('mean')),
                _v(f.get('std')),
                _v(f.get('median')),
                _v(f.get('min')),
                _v(f.get('max')),
                f.get('leakage', 'Low'),
            ])
        table = _data_table(doc,
            ['Feature Name', 'Var Type', 'Unique', 'Missing', 'Mean', 'Std Dev',
             'Median', 'Min', 'Max', 'Target Leakage'],
            fs_rows,
            [1.5, 0.7, 0.55, 0.6, 0.65, 0.65, 0.65, 0.65, 0.65, 0.85])
        # Colour leakage cells
        for ri, f in enumerate(feature_stats):
            leakage = f.get('leakage', 'Low')
            cell = table.rows[ri + 1].cells[9]
            if cell.paragraphs[0].runs:
                if leakage == 'High':
                    cell.paragraphs[0].runs[0].font.color.rgb = COL_RED
                    cell.paragraphs[0].runs[0].font.bold = True
                elif leakage == 'Moderate':
                    cell.paragraphs[0].runs[0].font.color.rgb = COL_ORANGE
    else:
        _body(doc, "Feature statistics not available for this model. Retrain to capture this data.")

    doc.add_page_break()

    # ── SECTION 6: PERFORMANCE AND STABILITY ─────────────────────────────────
    _h1(doc, "6   Model Performance and Stability")

    _h2(doc, "6.1   Model Validation Stability")
    _body(doc,
        "To find patterns in a dataset from which it can make predictions, an algorithm must first "
        "learn from a historical example — typically from a historical dataset. One common problem "
        "that arises in this process is overfitting, where the model learns the training data too well "
        "and cannot generalise to unseen data.")
    _body(doc,
        "Sentinel uses standard modelling techniques to validate model performance and ensure that "
        "overfitting does not occur. Sentinel uses a robust model k-fold cross-validation approach, "
        "in addition to an independent hold-out set:")
    for b in [
        "Sentinel sets aside 20% of the training data as a hold-out dataset. This dataset is used "
        "to verify that the final model performs well on data that was not seen during training.",
        "For further model validation, the remainder of the data is divided into 5 cross-validation "
        "partitions using Stratified K-Fold, preserving the target class ratio in each fold.",
    ]:
        _bullet(doc, b)

    _h3(doc, "6.1.1   Cross Validation Scores")
    _body(doc,
        "Sentinel calculates the Cross Validation scores for each of the training data partitions "
        "or folds. The AUC metric is used to calculate the score for each fold.")
    _caption(doc, "Table 4 — Cross Validation Scores by Fold")
    if cv_fold_scores:
        fold_rows = [[f"Fold {i+1}", _fmt_pct(s)] for i, s in enumerate(cv_fold_scores)]
        if cv_auc_mean is not None:
            fold_rows.append(['Mean', _fmt_pct(cv_auc_mean)])
        if cv_auc_std is not None:
            fold_rows.append(['Std Dev', f"\u00b1 {_fmt_pct(cv_auc_std)}"])
        table = _data_table(doc,
            ['Fold', 'Cross Validation Score (AUC)'],
            fold_rows,
            [2.0, 2.5])
        # Bold the mean/std rows
        for row in table.rows[-2:]:
            for cell in row.cells:
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.bold = True
    else:
        _body(doc, "Cross-validation fold scores not available. Retrain to capture this data.")

    _h3(doc, "6.1.2   Data Partitioning Methodology")
    _body(doc,
        "Because the distribution of the target in a binary classification project may be imbalanced, "
        "the modelling partitions were randomly selected using a stratified sampling approach. "
        "Stratified sampling ensures that the ratio of positive-to-negative target observations is "
        "preserved in both the training and test partitions, and in each cross-validation fold.")
    _kv_table(doc, [
        ('Train / Test Split',    '80% training / 20% hold-out (stratified)'),
        ('Cross-Validation',      '5-Fold Stratified K-Fold on training partition'),
        ('Stratification',        'target variable (preserves class ratio in all folds)'),
        ('Random Seed',           '42 (reproducible)'),
        ('Missing Value Strategy','Median imputation — fitted on training set only (no leakage)'),
        ('Class Imbalance',       'Balanced class weights (LR/RF) / scale_pos_weight=neg÷pos (XGBoost)'),
    ])

    _spacer(doc)
    _h2(doc, "6.2   Model Performance (Sample Scores)")
    _body(doc,
        "As an additional layer of model validity, Sentinel evaluated the model on the independent "
        "hold-out test set, which was not seen during training or cross-validation. "
        "The performance results are presented below:")
    _caption(doc, "Table 5 — Model Performance Scores")
    perf_rows = []
    if cv_auc_mean is not None:
        perf_rows.append(['cross_validation', _fmt_pct(cv_auc_mean)])
    perf_rows.append(['holdout', _fmt_pct(auc)])
    if cv_auc_std is not None:
        perf_rows.append(['cv_std_dev', f"\u00b1 {_fmt_pct(cv_auc_std)}"])
    perf_rows.append(['gini_coefficient', _fmt_pct(gini)])
    _data_table(doc,
        ['Scoring Type', 'AUC Score'],
        perf_rows,
        [2.5, 2.5])

    _spacer(doc)
    _h2(doc, "6.3   Sensitivity Testing and Analysis")

    _h3(doc, "6.3.1   Lift Chart")
    _body(doc,
        "The Lift Chart sorts and groups scores into equal-sized bins (deciles), depicting how well "
        "the model segments the target population. Each decile represents 10% of the scored population. "
        "A well-performing model produces a monotonically increasing pattern of default rates from the "
        "lowest-risk decile to the highest-risk decile.")
    _caption(doc, "Table 6 — Actual Default Rate by Score Decile (Decile 1 = lowest predicted risk)")

    if calibration:
        total   = sum(d['count'] for d in calibration)
        avg_bad = (sum(d['actual_rate'] * d['count'] for d in calibration) / total) if total > 0 else 0
        cal_rows    = []
        above_flags = []
        for d in calibration:
            lift  = d['actual_rate'] / avg_bad if avg_bad > 0 else 0
            above = d['actual_rate'] > avg_bad
            above_flags.append(above)
            cal_rows.append([
                str(d['decile']),
                f"{d.get('min_score', 0):.3f} \u2013 {d.get('max_score', 0):.3f}",
                f"{d['count']:,}",
                f"{d['actual_rate'] * 100:.2f}%",
                f"{lift:.2f}\u00d7",
                '\u25b2 Above avg' if above else '\u25bc Below avg',
            ])
        cal_rows.append(['Avg', '—', f"{total:,}", f"{avg_bad * 100:.2f}%", '1.00\u00d7', '—'])
        table = _data_table(doc,
            ['Decile', 'Score Range', 'Count', 'Default Rate', 'Lift Ratio', 'vs. Average'],
            cal_rows,
            [0.6, 1.6, 0.75, 1.0, 0.85, 0.95])
        for ri, above in enumerate(above_flags):
            cell = table.rows[ri + 1].cells[3]
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.color.rgb = COL_RED if above else COL_GREEN
        for cell in table.rows[-1].cells:
            _set_cell_bg(cell, HEX_SLATE)
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.color.rgb = COL_WHITE
                cell.paragraphs[0].runs[0].font.bold = True
    else:
        _body(doc, "Decile calibration data not available for this model.")

    _spacer(doc)
    _h3(doc, "6.3.2   Key Relationships")
    _body(doc,
        "Feature Impact measures the contribution of each feature to the model's predictions by "
        "observing how the model's output changes when each feature is permuted. The normalised "
        "impact score is scaled relative to the most important feature (score = 1.0).")
    _caption(doc, "Table 7 — Feature Impact (Top Features by Importance)")
    if feature_importance:
        max_imp = max((f.get('importance', 0) for f in feature_importance), default=1) or 1
        fi_rows = []
        for feat in feature_importance:
            imp = feat.get('importance', 0)
            norm = imp / max_imp
            fi_rows.append([
                feat.get('feature', '—'),
                f"{norm:.4f}",
                f"{imp:.4f}",
                feat.get('impact', 'Variable'),
            ])
        table = _data_table(doc,
            ['Feature Name', 'Impact Normalised', 'Impact Unnormalised', 'Direction'],
            fi_rows,
            [2.3, 1.3, 1.4, 1.3])
        for ri, feat in enumerate(feature_importance):
            impact = feat.get('impact', '')
            cell = table.rows[ri + 1].cells[3]
            if cell.paragraphs[0].runs:
                if 'Increases' in impact:
                    cell.paragraphs[0].runs[0].font.color.rgb = COL_RED
                elif 'Decreases' in impact:
                    cell.paragraphs[0].runs[0].font.color.rgb = COL_GREEN
    else:
        _body(doc, "Feature importance data not available for this model.")

    _spacer(doc)
    _h3(doc, "6.3.3   Sensitivity Analysis")
    _user_fill(doc,
        "In the case of linear models, users can gain considerable insight into the structure and "
        "interpretation of the model by examining its coefficients. For non-linear models, partial "
        "dependence plots or SHAP values should be used. Document the sensitivity analysis performed "
        "for this model, including any stress testing against out-of-time or out-of-population samples.")

    _spacer(doc)
    _h3(doc, "6.3.4   Accuracy (Receiver Operating Characteristic)")
    _body(doc,
        "A confusion matrix reports true versus predicted values at a classification threshold of 0.5. "
        "The ROC curve allows exploration of the trade-off between true positive rate (sensitivity) "
        "and false positive rate (1 - specificity) across all thresholds.")
    _body(doc, "The confusion matrix statistics for this model are reported below:")
    _caption(doc, "Table 8 — Classification Statistics at Threshold = 0.50")
    if classification_metrics:
        cm_rows = [[
            _fmt_pct(classification_metrics.get('f1')),
            _fmt_pct(classification_metrics.get('tpr')),
            _fmt_pct(classification_metrics.get('fpr')),
            _fmt_pct(classification_metrics.get('tnr')),
            _fmt_pct(classification_metrics.get('ppv')),
            _fmt_pct(classification_metrics.get('npv')),
            _fmt_pct(classification_metrics.get('accuracy')),
            _fmt_pct(classification_metrics.get('mcc')),
        ]]
        _data_table(doc,
            ['F1 Score', 'True Positive\nRate', 'False Positive\nRate', 'True Negative\nRate',
             'Positive Pred.\nValue', 'Negative Pred.\nValue', 'Accuracy', 'Matthews\nCorr. Coeff.'],
            cm_rows,
            [0.75, 0.85, 0.85, 0.85, 0.85, 0.85, 0.75, 0.85])
        _spacer(doc)
        _body(doc, "Where the values reported are:")
        for label, desc in [
            ("F1 Score",                    "A measure of the model's accuracy computed from precision and recall."),
            ("True Positive Rate (TPR)",     "Sensitivity / recall. The ratio of true positives to all actual positives."),
            ("False Positive Rate (FPR)",    "Fallout. The ratio of false positives to all actual negatives."),
            ("True Negative Rate (TNR)",     "Specificity. The ratio of true negatives to all actual negatives."),
            ("Positive Predictive Value",    "Precision. Of all positive predictions, the percentage that were correct."),
            ("Negative Predictive Value",    "Of all negative predictions, the percentage that were correct."),
            ("Accuracy",                     "The percentage of correctly classified instances overall."),
            ("Matthews Correlation Coeff.",  "Measure of model quality when classes are of very different sizes (imbalanced)."),
        ]:
            p = doc.add_paragraph(style='List Bullet')
            r1 = p.add_run(f"{label}: ")
            r1.font.bold = True
            r1.font.size = Pt(9)
            r2 = p.add_run(desc)
            r2.font.size = Pt(9)
    else:
        _body(doc, "Classification metrics not available for this model. Retrain to capture this data.")

    doc.add_page_break()

    # ── SECTION 7: IMPLEMENTATION ─────────────────────────────────────────────
    _h1(doc, "7   Model Implementation and Output Reporting")

    _h2(doc, "7.1   Version Control")
    _body(doc,
        "Sentinel handles model and project version control automatically by assigning each model a "
        "unique immutable Model ID at creation time. The Model ID represents the version of the model "
        "that was trained on a specific dataset and can be used to retrieve or audit the serialised "
        "model artifact at any time.")
    _body(doc,
        "The serialised model artifact is stored in object storage at the path below and can be "
        "reloaded using joblib for audit, re-scoring, or challenger model comparison.")
    _data_table(doc,
        ['Field', 'Value'],
        [
            ['Model ID',                model.id or '—'],
            ['Model Name',              model.name or '—'],
            ['Decision System ID',      model.decision_system_id or '—'],
            ['Algorithm',               algo_display],
            ['Artifact Path',           model.artifact_path or '—'],
            ['Serialisation Format',    'joblib (scikit-learn compatible)'],
            ['Training Framework',      'Sentinel v1.0 · scikit-learn 1.4 · XGBoost 2.0'],
            ['Documentation Generated', now_str],
        ],
        [2.2, 4.3])

    _spacer(doc)
    _h2(doc, "7.2   Monitoring and Ongoing Validation")
    _user_fill(doc,
        "Define the monitoring schedule, key performance indicators, and the performance thresholds "
        "that will trigger a model review or redevelopment. Include the process for periodic "
        "re-validation as required by SR-11-7 and your organisation's Model Risk Management policy. "
        "Typical monitoring metrics include: population stability index (PSI), Gini coefficient "
        "drift, and default rate deviation from model prediction.")

    _spacer(doc)
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run(
        f"Generated by Sentinel AI \u00b7 {now_str} \u00b7 Model ID: {model.id} \u00b7 "
        "Sections marked [REQUIRED: User Input] require completion by the Model Owner.")
    fr.font.size = Pt(7)
    fr.font.italic = True
    fr.font.color.rgb = COL_MUTED

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
